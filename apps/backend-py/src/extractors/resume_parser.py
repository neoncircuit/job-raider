"""
Job Raider - Resume Parser

This module provides functionality to parse resumes from PDF and DOCX files
and extract structured user profile information.

Author: Job Raider
Date: 2026-04-20
"""

import logging
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from pypdf import PdfReader

from ..llm.base import Message, MessageType
from ..llm.router import LLMRouter, TaskType
from ..models.job_listing import ExperienceLevel
from ..models.user_profile import (
    ApprenticeshipContract,
    Certification,
    ContactInfo,
    Education,
    Project,
    Skill,
    SkillCategory,
    TargetJob,
    UserProfile,
    WorkExperience,
)

logger = logging.getLogger(__name__)

# Resume section headers that end a Skills / Technical Skills block.
_NEXT_SECTION_HEADER = re.compile(
    r"(?im)^(?:work\s+experience|professional\s+experience|employment|"
    r"experience|projects?|education|certifications?|awards?|"
    r"publications?|summary|objective|profile|about)\s*:?\s*$"
)

# Skills section headers (standalone line or colon-led inline list).
# Does not match titles like "Skills Breakdown" or "Skills Radar".
_SKILLS_SECTION_HEADER = re.compile(
    r"(?im)^(?P<header>technical\s+skills|tech\s+skills|skills)\s*"
    r"(?::\s*(?P<inline>.*)?)?$"
)

# Projects section header (standalone).
_PROJECTS_SECTION_HEADER = re.compile(r"(?im)^projects?\s*:?\s*$")

# Category labels sometimes embedded inside a skills block.
_SKILL_CATEGORY_PREFIX = re.compile(
    r"^(?:programming\s+languages?|languages?|frameworks?|tools?|"
    r"databases?|cloud|domains?|other|methodologies?)\s*:\s*",
    re.IGNORECASE,
)

# Project highlight / bullet markers — never treated as tech-stack lines.
_PROJECT_BULLET_PREFIX = re.compile(r"^[-*•·▪◦‣⁃]\s+")

# Sentence-like openers that appear in prose bullets, not tool names.
_PROSE_TOKEN_STARTERS = frozenset(
    {
        "built",
        "building",
        "combining",
        "keeping",
        "designed",
        "implemented",
        "developed",
        "created",
        "using",
        "worked",
        "led",
        "leading",
        "improved",
        "reduced",
        "achieved",
        "responsible",
        "focused",
        "helping",
        "enabled",
        "delivered",
        "maintained",
        "integrated",
        "automated",
        "optimized",
    }
)

# LLM-invented umbrella labels — keep only if the phrase appears in resume text.
_INVENTED_UMBRELLA_SKILLS = frozenset(
    {
        "machine learning",
        "web development",
        "deep learning",
        "data science",
        "ai",
        "artificial intelligence",
    }
)

_PROGRAMMING_LANGUAGE_NAMES = frozenset(
    {
        "python",
        "javascript",
        "typescript",
        "java",
        "c++",
        "c#",
        "go",
        "golang",
        "rust",
        "ruby",
        "php",
        "swift",
        "kotlin",
        "scala",
        "r",
        "matlab",
        "sql",
        "html",
        "css",
    }
)

_FRAMEWORK_NAMES = frozenset(
    {
        "react",
        "angular",
        "vue",
        "next.js",
        "nextjs",
        "nuxt",
        "django",
        "flask",
        "fastapi",
        "spring",
        "express",
        "express.js",
        "node.js",
        "nodejs",
        ".net",
        "convex",
        "inngest",
        "letta",
    }
)

_DATABASE_NAMES = frozenset(
    {
        "postgresql",
        "postgres",
        "mysql",
        "mongodb",
        "redis",
        "sqlite",
        "elasticsearch",
        "dynamodb",
        "cassandra",
        "neo4j",
    }
)

_CLOUD_NAMES = frozenset(
    {
        "aws",
        "azure",
        "gcp",
        "google cloud",
        "azure openai",
        "google gen ai",
    }
)

_TOOL_NAMES = frozenset(
    {
        "git",
        "github",
        "gitlab",
        "docker",
        "kubernetes",
        "terraform",
        "ansible",
        "jenkins",
        "ci/cd",
        "cicd",
        "ollama",
        "qwen2.5",
        "qwen",
    }
)


class ResumeParser:
    """
    Parse resumes from PDF and DOCX files.

    Extracts structured user profile information including
    skills, experience, projects, and education.
    """

    def __init__(self, llm_router: Optional[LLMRouter] = None):
        """
        Initialize the resume parser.

        Args:
            llm_router: Optional LLM router for AI-based parsing
        """
        self.llm_router = llm_router

    def parse_file(self, file_path: str) -> UserProfile:
        """
        Parse a resume file (PDF or DOCX).

        Args:
            file_path: Path to resume file

        Returns:
            UserProfile with extracted information

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        # Determine file type and extract text
        if path.suffix.lower() == ".pdf":
            text = self._extract_text_from_pdf(path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            text = self._extract_text_from_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

        # Parse text into profile
        return self.parse_text(text)

    def parse_text(self, text: str) -> UserProfile:
        """
        Parse resume text into a UserProfile.

        Attaches ``metadata["resume_parse"]`` with parse datetime, duration,
        model, and method so the Profile UI can confirm a CV actually ran.

        Args:
            text: Resume text content

        Returns:
            UserProfile with extracted information
        """
        started = time.perf_counter()
        parsed_at = datetime.now()
        if self.llm_router:
            profile = self._parse_with_llm(text)
        else:
            profile = self._parse_rule_based(text)
            self._set_resume_parse_meta(
                profile,
                method="rule_based",
                model="rule-based",
                provider=None,
            )

        duration_ms = int((time.perf_counter() - started) * 1000)
        meta = profile.metadata.setdefault("resume_parse", {})
        meta["parsed_at"] = parsed_at.isoformat()
        meta["duration_ms"] = duration_ms
        return profile

    def _set_resume_parse_meta(
        self,
        profile: UserProfile,
        *,
        method: str,
        model: Optional[str],
        provider: Optional[str],
    ) -> None:
        """
        Record how a resume parse was produced on the profile metadata.

        Args:
            profile: Profile to update in place.
            method: ``llm`` or ``rule_based``.
            model: Model id used, or ``rule-based`` when no LLM ran.
            provider: Provider name when known (e.g. ``ollama``).
        """
        if profile.metadata is None:
            profile.metadata = {}
        existing = profile.metadata.get("resume_parse")
        meta: Dict[str, Any] = existing if isinstance(existing, dict) else {}
        meta["method"] = method
        meta["model"] = model or "unknown"
        if provider:
            meta["provider"] = provider
        elif "provider" not in meta:
            meta["provider"] = None
        profile.metadata["resume_parse"] = meta

    def _resume_parse_route_identity(
        self, model: Optional[str]
    ) -> tuple[Optional[str], Optional[str]]:
        """
        Resolve provider/model labels for a resume-parsing LLM response.

        Args:
            model: Model id reported on ``LLMResponse``.

        Returns:
            ``(provider, model)`` using the RESUME_PARSING route when possible.
        """
        if not self.llm_router:
            return None, model
        route = self.llm_router.routes.get(TaskType.RESUME_PARSING)
        if route is None:
            return None, model
        if model and route.fallback_model and model == route.fallback_model:
            return route.fallback_provider, model
        return route.primary_provider, model or route.primary_model

    def _extract_text_from_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(str(path))
            text = ""

            for page in reader.pages:
                text += page.extract_text() + "\n"

            return text

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_text_from_docx(self, path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(str(path))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def _parse_with_llm(self, text: str) -> UserProfile:
        """
        Parse resume using LLM.

        Args:
            text: Resume text content

        Returns:
            UserProfile with extracted information
        """
        # Keep the full text for deterministic Technical Skills extraction;
        # only truncate what is sent to the LLM.
        full_text = text
        max_length = 8000  # For models with 32K context
        llm_text = text[:max_length] if len(text) > max_length else text

        # Prepare prompt
        system_prompt = """You are a resume parser. Extract structured information from resume documents.
Be thorough and capture all relevant details. Return information in valid JSON format."""

        user_prompt = f"""Parse this resume and extract structured information:

RESUME TEXT:
{llm_text}

Return a JSON object with these exact fields:
{{
  "basics": {{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State/Country - this must be a physical geographic location (e.g. 'Singapore, Singapore', 'New York, NY', 'London, UK', 'Tokyo, Japan'), NOT a job title, role, or department name",
    "linkedin": "LinkedIn URL if present",
    "github": "GitHub URL if present"
  }},
  "summary": "Professional summary",
  "skills": {{
    "programming_languages": ["every programming language named on the resume"],
    "frameworks": ["every framework/library named on the resume"],
    "tools": ["every tool named on the resume"],
    "databases": ["every database named on the resume"],
    "cloud": ["every cloud platform named on the resume"],
    "domains": ["only domain phrases that appear verbatim on the resume — do not invent"]
  }},
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "location": "Location",
      "start_date": "YYYY-MM",
      "end_date": "YYYY-MM or 'present'",
      "highlights": ["bullet1", "bullet2", ...]
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "description": "Description",
      "technologies": ["tech1", "tech2"],
      "url": "URL if present",
      "highlights": ["achievement1", "achievement2"]
    }}
  ],
  "education": [
    {{
      "degree": "Degree Name",
      "school": "School Name",
      "location": "Location",
      "graduation_year": "YYYY",
      "gpa": "GPA if present"
    }}
  ],
  "certifications": [
    {{"name": "Certification", "issuer": "Issuer", "date": "YYYY-MM"}}
  ],
  "target_job": {{
    "keywords": ["target job titles or roles inferred from the candidate's profile. For experienced candidates, use work history. For recent graduates or students, infer from degree major, coursework, projects, and skills instead. e.g. Software Engineer, Data Analyst, Research Assistant"],
    "locations": ["preferred work locations inferred from current location or stated preferences"],
    "experience_levels": ["inferred career stage: one or more of Entry Level, Mid Level, Senior, Lead, Principal, Executive, Internship. Use Entry Level or Internship for recent graduates with little or no work experience."],
    "remote_preference": true or false based on any stated remote work preference,
    "industries": ["target industries inferred from the candidate's background. For experienced candidates, use work history. For recent graduates, infer from degree field, projects, and skills instead. e.g. Technology, Finance, Healthcare"]
  }},
  "apprenticeship": {{
    "field": "required field of work if the candidate has an apprenticeship/traineeship contract obligation, e.g. 'AI/ML', 'Healthcare'. Leave as null if no such contract exists.",
    "duration_months": "remaining contract duration in months as integer, or null",
    "employer": "sponsoring employer or organization name, or null",
    "start_date": "YYYY-MM or null",
    "end_date": "YYYY-MM or null",
    "is_active": true or false
  }}
}}

IMPORTANT: Return ONLY valid JSON. No markdown, no explanations, just the JSON object.
IMPORTANT: The "location" in "basics" must be a real city and state/country (e.g. "Singapore, Singapore" or "New York, NY" or "London, UK"), never a job title or role. City-states like Singapore are valid locations.
IMPORTANT: If the candidate has no work experience (recent graduate or student), still populate target_job keywords and industries based on their education, projects, and skills. Default experience_levels to ["Entry Level"] or ["Internship"] as appropriate.
IMPORTANT: If the resume mentions an apprenticeship, traineeship, or sponsored training program with a work obligation, populate the "apprenticeship" object. If no such program is mentioned, set all apprenticeship fields to null.
IMPORTANT: For skills, extract EVERY named skill under Technical Skills / Skills verbatim. Do not summarize into a short representative subset. Do not invent umbrella domains such as "Machine Learning" or "Web Development" unless those exact phrases appear on the resume. Atomic skills such as RAG, Prompt Engineering, Vector Databases, FastAPI, Convex, TypeScript, and SQL must be kept as separate entries when present. A flat skills array is also acceptable.
IMPORTANT: For each project, technologies must come only from that project's own tech/stack line in the Projects section. Do not copy technologies from Experience, other projects, or the global Technical Skills list into a project. Include every tech on that project's line — do not summarize a long stack into a short subset."""

        messages = [
            Message(role=MessageType.SYSTEM, content=system_prompt),
            Message(role=MessageType.USER, content=user_prompt),
        ]

        # Rule-based wipe only when the LLM call or JSON extract/load fails.
        # Mapping edge cases are handled inside _create_profile_from_dict so a
        # null array item / bad URL / bad GPA cannot empty experience/projects.
        response_model: Optional[str] = None
        try:
            response = self.llm_router.generate(
                messages=messages,
                task_type=TaskType.RESUME_PARSING,
                temperature=0.3,
            )
            response_model = response.model
            json_match = re.search(r"\{.*\}", response.content, re.DOTALL)
            if not json_match:
                raise ValueError("Failed to extract JSON from LLM response")

            import json

            data = json.loads(json_match.group(0))
            if not isinstance(data, dict):
                raise ValueError("LLM resume JSON root must be an object")
        except Exception as e:
            logger.error(f"LLM resume parsing failed, falling back to rule-based: {e}")
            profile = self._parse_rule_based(full_text)
            self._set_resume_parse_meta(
                profile,
                method="rule_based",
                model="rule-based",
                provider=None,
            )
            return profile

        # Technical Skills section is the sole SoT for profile.skills when
        # present. Project/experience stacks stay on their own fields;
        # project technologies are overwritten from each project's own block.
        try:
            profile = self._create_profile_from_dict(data)
        except Exception as e:
            logger.error(
                "Resume JSON mapped with failures; keeping partial profile: %s",
                e,
            )
            profile = self._create_profile_from_dict({})

        section_skills = self._skills_from_technical_section(full_text)
        profile.skills = self._merge_skills(
            section_skills,
            list(profile.skills or []),
            full_text,
        )
        self._apply_section_project_technologies(profile, full_text)
        provider, model = self._resume_parse_route_identity(response_model)
        self._set_resume_parse_meta(
            profile,
            method="llm",
            model=model,
            provider=provider,
        )
        return profile

    def _as_dict(self, value: Any) -> Dict[str, Any]:
        """
        Coerce a value to a dict, treating null/non-dicts as empty.

        Args:
            value: Candidate mapping from LLM JSON.

        Returns:
            ``value`` when it is a dict, otherwise ``{}``.
        """
        return value if isinstance(value, dict) else {}

    def _iter_dict_items(self, value: Any) -> List[Dict[str, Any]]:
        """
        Yield only dict entries from a list-like LLM field.

        Args:
            value: Candidate list from LLM JSON (may contain nulls).

        Returns:
            List of dict items; non-dicts are skipped.
        """
        if not isinstance(value, list):
            return []
        return [item for item in value if isinstance(item, dict)]

    def _experience_entries(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Resolve experience entries from primary or aliased LLM keys.

        Args:
            data: Top-level parsed resume JSON.

        Returns:
            Dict entries under ``experience``, ``work_experience``, or ``work``.
        """
        for key in ("experience", "work_experience", "work"):
            entries = self._iter_dict_items(data.get(key))
            if entries:
                return entries
        return []

    def _coerce_gpa(self, value: Any) -> Optional[float]:
        """
        Coerce a GPA value to a float in ``[0, 4]``, else drop it.

        Args:
            value: Raw GPA from LLM JSON (number, string, or junk).

        Returns:
            Valid GPA float, or None when missing/unusable.
        """
        if value is None or value == "":
            return None
        try:
            gpa = float(str(value).strip())
        except (TypeError, ValueError):
            return None
        if gpa < 0.0 or gpa > 4.0:
            return None
        return gpa

    def _create_profile_from_dict(self, data: Dict[str, Any]) -> UserProfile:
        """
        Create UserProfile from extracted dictionary data.

        Null/non-dict nested objects and list items are skipped or coerced so
        one bad LLM field cannot wipe Experience/Projects for the whole upload.

        Args:
            data: Parsed resume JSON object (may be empty).

        Returns:
            UserProfile with as many valid entries as could be mapped.
        """
        if not isinstance(data, dict):
            data = {}

        # Normalize top-level keys: LLM may return different field names
        basics = self._as_dict(data.get("basics") or data.get("basic_info"))
        summary = (
            data.get("summary")
            or data.get("profile")
            or data.get("professional_summary")
            or ""
        )
        if not isinstance(summary, str):
            summary = str(summary) if summary else ""
        skills_dict = data.get("skills") or data.get("technical_skills") or {}
        # If skills is a flat list instead of categorized dict, convert it
        if isinstance(skills_dict, list):
            skills_dict = {"programming_languages": skills_dict}
        elif not isinstance(skills_dict, dict):
            skills_dict = {}

        email = basics.get("email") or "user@example.com"
        try:
            contact = ContactInfo(
                email=str(email),
                phone=basics.get("phone"),
                location=basics.get("location") or "Unknown",
                linkedin=self._normalize_url(basics.get("linkedin")),
                github=self._normalize_url(basics.get("github")),
            )
        except Exception as e:
            logger.warning("Invalid contact fields; using defaults: %s", e)
            contact = ContactInfo(email="user@example.com", location="Unknown")

        # Parse skills (using normalized skills_dict from above)
        skills = []

        skill_category_map = {
            "programming_languages": SkillCategory.PROGRAMMING_LANGUAGE,
            "programming_language": SkillCategory.PROGRAMMING_LANGUAGE,
            "frameworks": SkillCategory.FRAMEWORK,
            "framework": SkillCategory.FRAMEWORK,
            "tools": SkillCategory.TOOL,
            "tool": SkillCategory.TOOL,
            "cloud": SkillCategory.CLOUD,
            "database": SkillCategory.DATABASE,
            "databases": SkillCategory.DATABASE,
            "domains": SkillCategory.DOMAIN,
            "domain": SkillCategory.DOMAIN,
            "languages": SkillCategory.LANGUAGE,
            "language": SkillCategory.LANGUAGE,
        }

        for category_str, skill_list in skills_dict.items():
            if not isinstance(skill_list, list):
                continue
            category = skill_category_map.get(
                str(category_str).strip().lower(), SkillCategory.OTHER
            )
            for skill_name in skill_list:
                if not skill_name:
                    continue
                try:
                    skills.append(
                        Skill(name=str(skill_name).strip(), category=category)
                    )
                except Exception as e:
                    logger.warning("Skipping invalid skill %r: %s", skill_name, e)

        # Parse experience
        experience = []
        for exp_data in self._experience_entries(data):
            try:
                title = exp_data.get("title") or exp_data.get("position") or "Unknown"
                company = (
                    exp_data.get("company") or exp_data.get("organization") or "Unknown"
                )
                highlights = (
                    exp_data.get("highlights")
                    or exp_data.get("achievements")
                    or exp_data.get("responsibilities")
                    or []
                )

                # Prefer explicit start_date/end_date; fall back to "dates" string
                start_date = self._parse_date(exp_data.get("start_date"))
                end_date_str = exp_data.get("end_date")
                if not start_date and not end_date_str:
                    dates_raw = (
                        exp_data.get("dates") or exp_data.get("date_range") or ""
                    )
                    start_date, end_date_str = self._split_date_range(dates_raw)
                end_date = (
                    None
                    if (str(end_date_str or "")).lower() in ("present", "current", "")
                    else self._parse_date(end_date_str)
                )

                description = exp_data.get("description") or ""
                if isinstance(description, list):
                    description = " ".join(str(h) for h in description)
                if not isinstance(highlights, list):
                    highlights = [str(highlights)] if highlights else []
                else:
                    highlights = [str(item) for item in highlights if item]
                description, highlights = self._dedupe_description_and_highlights(
                    str(description) if description else "",
                    highlights,
                )
                technologies = exp_data.get("technologies") or []
                if not isinstance(technologies, list):
                    technologies = [str(technologies)] if technologies else []
                else:
                    technologies = [str(item) for item in technologies if item]

                experience.append(
                    WorkExperience(
                        title=str(title),
                        company=str(company),
                        location=exp_data.get("location"),
                        start_date=start_date or datetime.now(),
                        end_date=end_date,
                        description=description[:500] if description else None,
                        highlights=highlights,
                        technologies=technologies,
                    )
                )
            except Exception as e:
                logger.warning("Skipping invalid experience entry: %s", e)

        # Parse projects
        projects = []
        for proj_data in self._iter_dict_items(data.get("projects")):
            try:
                description = proj_data.get("description") or ""
                highlights = (
                    proj_data.get("highlights") or proj_data.get("achievements") or []
                )
                if not isinstance(highlights, list):
                    highlights = [str(highlights)] if highlights else []
                else:
                    highlights = [str(item) for item in highlights if item]
                if isinstance(description, list):
                    description = " ".join(str(h) for h in description)
                description, highlights = self._dedupe_description_and_highlights(
                    str(description) if description else "",
                    highlights,
                )
                technologies = (
                    proj_data.get("technologies") or proj_data.get("tech_stack") or []
                )
                if not isinstance(technologies, list):
                    technologies = [str(technologies)] if technologies else []
                else:
                    technologies = [str(item) for item in technologies if item]
                if self._description_duplicates_technologies(description, technologies):
                    description = ""

                projects.append(
                    Project(
                        name=proj_data.get("name")
                        or proj_data.get("title")
                        or "Unknown",
                        description=description[:500] if description else "",
                        technologies=technologies,
                        url=self._normalize_url(proj_data.get("url")),
                        highlights=highlights,
                    )
                )
            except Exception as e:
                logger.warning("Skipping invalid project entry: %s", e)

        # Parse education
        education = []
        for edu_data in self._iter_dict_items(data.get("education")):
            try:
                grad_year = edu_data.get("graduation_year") or edu_data.get("year")
                end_date = None
                if grad_year:
                    try:
                        end_date = datetime(int(str(grad_year)[:4]), 6, 1)
                    except (ValueError, TypeError):
                        pass

                education.append(
                    Education(
                        degree=edu_data.get("degree")
                        or edu_data.get("qualification")
                        or "Unknown",
                        school=edu_data.get("school")
                        or edu_data.get("university")
                        or edu_data.get("institution")
                        or "Unknown",
                        location=edu_data.get("location"),
                        end_date=end_date,
                        gpa=self._coerce_gpa(edu_data.get("gpa")),
                    )
                )
            except Exception as e:
                logger.warning("Skipping invalid education entry: %s", e)

        # Parse certifications
        certifications = []
        for cert_data in self._iter_dict_items(data.get("certifications")):
            try:
                issue_date = self._parse_date(cert_data.get("date"))
                certifications.append(
                    Certification(
                        name=cert_data.get("name") or "Unknown",
                        issuer=cert_data.get("issuer") or "Unknown",
                        issue_date=issue_date,
                    )
                )
            except Exception as e:
                logger.warning("Skipping invalid certification entry: %s", e)

        target_job = self._build_target_job(
            self._as_dict(data.get("target_job")),
        )
        apprenticeship = self._build_apprenticeship(
            self._as_dict(data.get("apprenticeship")),
        )

        return UserProfile(
            name=basics.get("name") or "User",
            contact=contact,
            summary=summary or None,
            skills=skills,
            experience=experience,
            projects=projects,
            education=education,
            certifications=certifications,
            targets=target_job,
            apprenticeship=apprenticeship,
        )

    def _build_target_job(self, data: Dict[str, Any]) -> TargetJob:
        """Build a TargetJob from extracted dictionary data.

        Maps experience level strings to the ExperienceLevel enum,
        falling back to NOT_SPECIFIED for unrecognized values.

        Args:
            data: Dictionary with target job preference fields.

        Returns:
            TargetJob with extracted preferences.
        """
        data = self._as_dict(data)
        level_map = {
            "entry level": ExperienceLevel.ENTRY,
            "entry": ExperienceLevel.ENTRY,
            "mid level": ExperienceLevel.MID,
            "mid": ExperienceLevel.MID,
            "senior": ExperienceLevel.SENIOR,
            "lead": ExperienceLevel.LEAD,
            "principal": ExperienceLevel.PRINCIPAL,
            "executive": ExperienceLevel.EXECUTIVE,
            "internship": ExperienceLevel.INTERNSHIP,
        }

        experience_levels: List[ExperienceLevel] = []
        levels_raw = data.get("experience_levels") or []
        if not isinstance(levels_raw, list):
            levels_raw = [levels_raw] if levels_raw else []
        for level_str in levels_raw:
            if level_str is None:
                continue
            normalized = str(level_str).strip().lower()
            matched = level_map.get(normalized, ExperienceLevel.NOT_SPECIFIED)
            if matched not in experience_levels:
                experience_levels.append(matched)

        def _str_list(value: Any) -> List[str]:
            if not isinstance(value, list):
                return []
            return [str(item) for item in value if item]

        return TargetJob(
            keywords=_str_list(data.get("keywords")),
            locations=_str_list(data.get("locations")),
            experience_levels=experience_levels,
            remote_preference=bool(data.get("remote_preference", False)),
            industries=_str_list(data.get("industries")),
        )

    def _build_apprenticeship(
        self, data: Dict[str, Any]
    ) -> Optional[ApprenticeshipContract]:
        """Build an ApprenticeshipContract from extracted dictionary data.

        Returns None if no apprenticeship data is present or the field
        is empty/null.

        Args:
            data: Dictionary with apprenticeship contract fields.

        Returns:
            ApprenticeshipContract if data is valid, None otherwise.
        """
        data = self._as_dict(data)
        field = data.get("field")
        if not field:
            return None

        try:
            return ApprenticeshipContract(
                field=str(field),
                duration_months=data.get("duration_months"),
                employer=data.get("employer"),
                start_date=self._parse_date(data.get("start_date")),
                end_date=self._parse_date(data.get("end_date")),
                is_active=data.get("is_active", True),
            )
        except Exception as e:
            logger.warning("Skipping invalid apprenticeship data: %s", e)
            return None

    def _parse_rule_based(self, text: str) -> UserProfile:
        """
        Parse resume using rule-based extraction.

        This is a simplified version. In production, you'd want
        more sophisticated patterns or use a dedicated resume parsing library.

        Args:
            text: Resume text content

        Returns:
            UserProfile with extracted information
        """
        # Extract name (first line typically)
        lines = text.split("\n")
        name = lines[0].strip() if lines else "User"

        # Extract email
        email_match = re.search(r"\b[\w\.-]+@[\w\.-]+\.\w+\b", text)
        email = email_match.group(0) if email_match else "user@example.com"

        # Extract phone - more flexible patterns for international formats
        phone_patterns = [
            r"\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}",  # International
            r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",  # US format
            r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}",  # Simple format
            r"\b\d{10}\b",  # 10 digits
        ]
        phone = None
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                phone = phone_match.group(0)
                break

        # Extract location - improved pattern to avoid false matches
        # Look for common location indicators
        location = None

        # Blocklist of words that indicate a job title, not a location
        location_blocklist = {
            "assistant",
            "manager",
            "engineer",
            "developer",
            "analyst",
            "director",
            "coordinator",
            "specialist",
            "consultant",
            "intern",
            "associate",
            "designer",
            "administrator",
            "lead",
            "architect",
            "scientist",
            "officer",
            "executive",
            "supervisor",
            "technician",
        }

        # Try specific patterns first
        location_patterns = [
            r"(?:Location|City|Address|Based in?):\s*([A-Z][a-zA-Z\s]+?(?:,\s*[A-Z]{2}|,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)?(?:\n|$))",
            r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z]{2}|[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)",  # City, State/Country (e.g. "Singapore, Singapore")
        ]

        for pattern in location_patterns:
            location_match = re.search(pattern, text, re.IGNORECASE)
            if location_match:
                loc_text = location_match.group(1).strip()
                loc_words = loc_text.lower().split()
                # Validate: reasonable length and not a job title
                is_blocked = any(word in location_blocklist for word in loc_words)
                if 3 < len(loc_text) < 100 and not is_blocked:
                    location = loc_text
                    break

        # Fallback to unknown if no location found
        if not location:
            location = "Unknown"

        # Create basic profile
        section_skills = self._skills_from_technical_section(text)
        lexicon_skills = self._extract_skills_rule_based(text)
        return UserProfile(
            name=name,
            contact=ContactInfo(
                email=email,
                phone=phone,
                location=location,
            ),
            skills=self._merge_skills(section_skills, lexicon_skills, text),
            experience=[],
            projects=[],
            education=[],
        )

    def _normalize_prose_key(self, text: str) -> str:
        """
        Normalize free text for description/highlight equality checks.

        Args:
            text: Raw description or bullet string.

        Returns:
            Lowercased, whitespace-collapsed comparison key.
        """
        return re.sub(r"\s+", " ", (text or "").strip().lower())

    def _dedupe_description_and_highlights(
        self,
        description: str,
        highlights: List[str],
    ) -> tuple[str, List[str]]:
        """
        Ensure the summary line is not repeated as the first bullet.

        If ``description`` is empty and highlights exist, use the first
        highlight as the description and drop it from the bullet list. If the
        description already matches the first highlight verbatim (normalized),
        drop that highlight so the UI does not show the same sentence twice.

        Args:
            description: Role/project summary line (may be empty).
            highlights: Bullet achievements under that entry.

        Returns:
            Tuple of ``(description, highlights)`` with structural dedupe applied.
        """
        cleaned_highlights = [
            str(item).strip() for item in highlights if str(item).strip()
        ]
        cleaned_description = (description or "").strip()

        if not cleaned_description and cleaned_highlights:
            cleaned_description = cleaned_highlights[0]
            cleaned_highlights = cleaned_highlights[1:]
            return cleaned_description, cleaned_highlights

        if cleaned_description and cleaned_highlights:
            desc_key = self._normalize_prose_key(cleaned_description)
            while (
                cleaned_highlights
                and self._normalize_prose_key(cleaned_highlights[0]) == desc_key
            ):
                cleaned_highlights = cleaned_highlights[1:]

        return cleaned_description, cleaned_highlights

    def _normalize_skill_key(self, name: str) -> str:
        """
        Normalize a skill name for deduplication.

        Args:
            name: Raw skill string.

        Returns:
            Lowercased, whitespace-collapsed key.
        """
        return re.sub(r"\s+", " ", (name or "").strip().lower())

    def _extract_technical_skills_section(self, text: str) -> Optional[str]:
        """
        Locate the Technical Skills / Skills block in resume text.

        Args:
            text: Full resume text (not LLM-truncated).

        Returns:
            Section body text including any inline skills after the header,
            or None when no skills section is found.
        """
        if not text or not text.strip():
            return None

        lines = text.splitlines()
        for index, line in enumerate(lines):
            match = _SKILLS_SECTION_HEADER.match(line.strip())
            if not match:
                continue

            chunks: List[str] = []
            inline = (match.group("inline") or "").strip()
            if inline:
                chunks.append(inline)

            for following in lines[index + 1 :]:
                stripped = following.strip()
                if not stripped:
                    if chunks:
                        # Allow a single blank line inside the block.
                        chunks.append("")
                    continue
                if _NEXT_SECTION_HEADER.match(stripped):
                    break
                if _SKILLS_SECTION_HEADER.match(stripped):
                    break
                chunks.append(stripped)

            body = "\n".join(chunks).strip()
            if body:
                return body
        return None

    def _split_skill_tokens(self, section_text: str) -> List[str]:
        """
        Split a skills section into verbatim skill tokens.

        Preserves multi-word names and tokens that contain ``/`` such as
        ``CI/CD`` or ``REST APIs``.

        Args:
            section_text: Body of the Technical Skills section.

        Returns:
            Ordered list of non-empty skill names.
        """
        if not section_text:
            return []

        cleaned = re.sub(r"[•·▪◦‣⁃]\s*", ",", section_text)
        parts = re.split(r"[,|;]+|\n+", cleaned)
        tokens: List[str] = []
        seen: set[str] = set()
        for part in parts:
            token = _SKILL_CATEGORY_PREFIX.sub("", part.strip())
            token = token.strip(" -\t")
            if not token:
                continue
            key = self._normalize_skill_key(token)
            if key in seen:
                continue
            seen.add(key)
            tokens.append(token)
        return tokens

    def _categorize_skill_name(self, name: str) -> SkillCategory:
        """
        Assign a SkillCategory using a light name heuristic.

        Args:
            name: Skill display name.

        Returns:
            Best-effort category; methodology terms default to OTHER.
        """
        key = self._normalize_skill_key(name)
        if key in _PROGRAMMING_LANGUAGE_NAMES:
            return SkillCategory.PROGRAMMING_LANGUAGE
        if key in _FRAMEWORK_NAMES:
            return SkillCategory.FRAMEWORK
        if key in _DATABASE_NAMES:
            return SkillCategory.DATABASE
        if key in _CLOUD_NAMES:
            return SkillCategory.CLOUD
        if key in _TOOL_NAMES:
            return SkillCategory.TOOL
        return SkillCategory.OTHER

    def _skills_from_technical_section(self, text: str) -> List[Skill]:
        """
        Build Skill objects from a deterministic Technical Skills section parse.

        Args:
            text: Full resume text.

        Returns:
            Skills listed under Technical Skills / Skills, or an empty list.
        """
        section = self._extract_technical_skills_section(text)
        if not section:
            return []
        return [
            Skill(name=token, category=self._categorize_skill_name(token))
            for token in self._split_skill_tokens(section)
        ]

    def _is_tech_stack_token(self, token: str) -> bool:
        """
        Return whether a split token looks like a tool/tech name, not prose.

        Args:
            token: One comma-split fragment from a candidate tech line.

        Returns:
            True for short tool-like names (e.g. ``Docker``, ``Google Gen AI``);
            False for sentence clauses chopped on commas.
        """
        cleaned = (token or "").strip()
        if not cleaned:
            return False
        words = [word for word in re.split(r"\s+", cleaned) if word]
        if not words or len(words) > 4 or len(cleaned) > 48:
            return False
        first = words[0].lower().rstrip(".,;:")
        if first in _PROSE_TOKEN_STARTERS:
            return False
        padded = f" {cleaned.lower()} "
        if " the " in padded or " across " in padded or " with a " in padded:
            return False
        return True

    def _looks_like_bullet_or_prose(self, line: str) -> bool:
        """
        Return whether a line is project body prose (bullets / sentences).

        Args:
            line: Candidate Projects-section line.

        Returns:
            True when the line should never contribute project technologies.
        """
        stripped = (line or "").strip()
        if not stripped:
            return False
        if _PROJECT_BULLET_PREFIX.match(stripped):
            return True
        if self._looks_like_tech_line(stripped):
            return False
        words = stripped.split()
        if len(words) >= 8 and not re.search(r"[,|;]", stripped):
            return True
        if re.search(r"[,|;]", stripped):
            tokens = self._split_skill_tokens(stripped)
            if tokens and not all(self._is_tech_stack_token(token) for token in tokens):
                return True
        return False

    def _looks_like_tech_line(self, line: str) -> bool:
        """
        Return whether a line looks like a short comma-separated tech stack.

        Rejects bullet/prose lines even when they contain commas, so description
        sentences are not split into tag pills.

        Args:
            line: Candidate resume line under a project.

        Returns:
            True when every comma-split token is tool-like and there are >= 2.
        """
        stripped = (line or "").strip()
        if not stripped or _PROJECT_BULLET_PREFIX.match(stripped):
            return False
        if not re.search(r"[,|;]", stripped):
            return False
        tokens = self._split_skill_tokens(stripped)
        if len(tokens) < 2:
            return False
        return all(self._is_tech_stack_token(token) for token in tokens)

    def _extract_projects_section(self, text: str) -> Optional[str]:
        """
        Locate the Projects section body in resume text.

        Args:
            text: Full resume text.

        Returns:
            Projects section body, or None when not found.
        """
        if not text or not text.strip():
            return None

        lines = text.splitlines()
        for index, line in enumerate(lines):
            if not _PROJECTS_SECTION_HEADER.match(line.strip()):
                continue
            chunks: List[str] = []
            for following in lines[index + 1 :]:
                stripped = following.strip()
                if not stripped:
                    if chunks:
                        chunks.append("")
                    continue
                if _NEXT_SECTION_HEADER.match(
                    stripped
                ) and not _PROJECTS_SECTION_HEADER.match(stripped):
                    break
                chunks.append(stripped)
            body = "\n".join(chunks).strip()
            if body:
                return body
        return None

    def _normalize_project_key(self, name: str) -> str:
        """
        Normalize a project name for matching LLM output to section blocks.

        Strips affiliation suffixes in parentheses and collapses whitespace.

        Args:
            name: Project title as shown on the resume or returned by the LLM.

        Returns:
            Lowercased name key used for matching.
        """
        base = re.split(r"\s*\(", (name or "").strip(), maxsplit=1)[0]
        return re.sub(r"\s+", " ", base.lower()).strip()

    def _parse_inline_project_tech_line(
        self, line: str
    ) -> Optional[tuple[str, List[str]]]:
        """
        Parse a single line that contains both project name and tech stack.

        Handles forms such as ``Agent-C (AI Singapore) Python, Docker, ...``,
        ``Job Raider: Python, FastAPI, ...``, and
        ``Agent-C (AI Singapore) | Python, Docker, ...``.

        Args:
            line: Candidate Projects-section line.

        Returns:
            ``(project_name, techs)`` when both parts are present, else None.
        """
        stripped = (line or "").strip()
        if not stripped:
            return None

        pipe_match = re.match(
            r"^(?P<name>.+?)\s*\|\s*(?P<techs>.+)$",
            stripped,
        )
        if pipe_match and self._looks_like_tech_line(pipe_match.group("techs")):
            name = pipe_match.group("name").strip()
            techs = self._split_skill_tokens(pipe_match.group("techs"))
            if name and techs:
                return name, techs

        colon_match = re.match(
            r"^(?P<name>[^:]+?)\s*:\s*(?P<techs>.+)$",
            stripped,
        )
        if colon_match and self._looks_like_tech_line(colon_match.group("techs")):
            name = colon_match.group("name").strip()
            techs = self._split_skill_tokens(colon_match.group("techs"))
            if name and techs:
                return name, techs

        paren_match = re.match(
            r"^(?P<name>.+?\([^)]+\))\s+(?P<techs>.+)$",
            stripped,
        )
        if paren_match and self._looks_like_tech_line(paren_match.group("techs")):
            name = paren_match.group("name").strip()
            techs = self._split_skill_tokens(paren_match.group("techs"))
            if name and techs:
                return name, techs

        return None

    def _parse_project_tech_blocks(self, text: str) -> Dict[str, List[str]]:
        """
        Parse per-project technology lines from the Projects section only.

        Only short tech-stack lines become tags. Bullet/prose body text is
        ignored even when it contains commas. Blank lines between the name and
        tech line are allowed; wrapped tech lines are merged until prose begins.

        Args:
            text: Full resume text.

        Returns:
            Mapping of project name line -> ordered tech tokens for that block.
        """
        section = self._extract_projects_section(text)
        if not section:
            return {}

        blocks: Dict[str, List[str]] = {}
        current_name: Optional[str] = None
        accepting_techs = False

        def _extend_techs(name: str, techs: List[str]) -> None:
            existing = blocks.setdefault(name, [])
            seen = {self._normalize_skill_key(token) for token in existing}
            for token in techs:
                if not self._is_tech_stack_token(token):
                    continue
                key = self._normalize_skill_key(token)
                if key in seen:
                    continue
                seen.add(key)
                existing.append(token)

        for raw_line in section.splitlines():
            line = raw_line.strip()
            if not line:
                # Keep name/accepting state across blank lines so PDF layouts
                # like "Agent-C\\n\\nPython, Docker, ..." still bind techs.
                continue

            # Prefer inline name+tech forms before prose detection so lines
            # like "Job Raider: Python, FastAPI" are not discarded as sentences.
            inline = self._parse_inline_project_tech_line(line)
            if inline is not None:
                name, techs = inline
                current_name = name
                accepting_techs = True
                _extend_techs(name, techs)
                continue

            if self._looks_like_bullet_or_prose(line):
                # Prose/highlights end the tech-stack region for this project.
                accepting_techs = False
                continue

            if self._looks_like_tech_line(line):
                if current_name and accepting_techs:
                    _extend_techs(current_name, self._split_skill_tokens(line))
                continue

            # New project header / name line within the Projects section.
            current_name = line
            accepting_techs = True
            blocks.setdefault(current_name, [])

        return {name: techs for name, techs in blocks.items() if techs}

    def _find_section_technologies_for_project(
        self,
        project_name: str,
        blocks: Dict[str, List[str]],
    ) -> Optional[List[str]]:
        """
        Match an LLM project name to a section-parsed tech block.

        Args:
            project_name: Project name from the structured profile.
            blocks: Mapping from section project titles to tech tokens.

        Returns:
            Tech list for the best-matching block, or None if no match.
        """
        if not project_name or not blocks:
            return None

        key = self._normalize_project_key(project_name)
        exact = {
            self._normalize_project_key(name): techs for name, techs in blocks.items()
        }
        if key in exact:
            return list(exact[key])

        for block_name, techs in blocks.items():
            block_key = self._normalize_project_key(block_name)
            if (
                key.startswith(block_key)
                or block_key.startswith(key)
                or key in block_key
                or block_key in key
            ):
                return list(techs)
        return None

    def _description_duplicates_technologies(
        self,
        description: str,
        technologies: List[str],
    ) -> bool:
        """
        Return whether a project description is just its tech-stack line.

        Args:
            description: Project summary text.
            technologies: Parsed technology tags for the same project.

        Returns:
            True when description is the same stack shown as tags.
        """
        cleaned = (description or "").strip()
        techs = [str(item).strip() for item in technologies if str(item).strip()]
        if not cleaned or not techs:
            return False

        joined = ", ".join(techs)
        if self._normalize_prose_key(cleaned) == self._normalize_prose_key(joined):
            return True

        if not self._looks_like_tech_line(cleaned):
            return False

        desc_keys = sorted(
            self._normalize_skill_key(token)
            for token in self._split_skill_tokens(cleaned)
        )
        tech_keys = sorted(self._normalize_skill_key(token) for token in techs)
        return desc_keys == tech_keys

    def _apply_section_project_technologies(
        self,
        profile: UserProfile,
        resume_text: str,
    ) -> None:
        """
        Overwrite project technologies from section-bound tech lines when present.

        This corrects LLM misattribution (e.g. Experience PostgreSQL onto Agent-C)
        and under-inclusion on long project stacks. Projects without a matching
        section tech line keep their LLM technologies. If ``description`` is only
        the same tech-stack line, clear it so tags are not duplicated as prose.

        Args:
            profile: Profile whose ``projects`` list may be updated in place.
            resume_text: Full resume text used for section parsing.
        """
        blocks = self._parse_project_tech_blocks(resume_text)
        for project in profile.projects or []:
            if blocks:
                techs = self._find_section_technologies_for_project(
                    project.name, blocks
                )
                if techs is not None:
                    project.technologies = techs
            if self._description_duplicates_technologies(
                project.description or "",
                list(project.technologies or []),
            ):
                project.description = ""

    def _resume_contains_phrase(self, resume_text: str, phrase: str) -> bool:
        """
        Return whether ``phrase`` appears in resume text as a whole token.

        Args:
            resume_text: Full resume body.
            phrase: Candidate skill or domain phrase.

        Returns:
            True when the phrase is present with token boundaries.
        """
        if not resume_text or not phrase:
            return False
        pattern = re.compile(
            rf"(?<![A-Za-z0-9]){re.escape(phrase.strip())}(?![A-Za-z0-9])",
            re.IGNORECASE,
        )
        return bool(pattern.search(resume_text))

    def _merge_skills(
        self,
        section_skills: List[Skill],
        llm_skills: List[Skill],
        resume_text: str,
    ) -> List[Skill]:
        """
        Resolve ``profile.skills`` with an explicit source-of-truth rule.

        When a Technical Skills / Skills section is present, that section is
        the sole source of ``profile.skills`` (verbatim, deduplicated). Project
        and experience tech stacks remain on those entities and are not folded
        into the skills count. When no section is found, fall back to
        LLM/lexicon skills with umbrella-domain filtering.

        Args:
            section_skills: Skills from the Technical Skills section.
            llm_skills: Skills from the LLM or rule-based lexicon.
            resume_text: Full resume text used to validate umbrella labels.

        Returns:
            Deduplicated skill list for ``profile.skills``.
        """
        merged: List[Skill] = []
        seen: set[str] = set()

        def _append(skill: Skill) -> None:
            if not skill.name:
                return
            key = self._normalize_skill_key(skill.name)
            if key in seen:
                return
            seen.add(key)
            merged.append(skill)

        if section_skills:
            for skill in section_skills:
                _append(skill)
            return merged

        for skill in llm_skills:
            key = self._normalize_skill_key(skill.name)
            if key in _INVENTED_UMBRELLA_SKILLS and not self._resume_contains_phrase(
                resume_text, skill.name
            ):
                continue
            _append(skill)

        return merged

    def _extract_skills_rule_based(self, text: str) -> List[Skill]:
        """
        Extract skills using rule-based lexicon patterns.

        Args:
            text: Resume text content.

        Returns:
            Skills matched by the fixed regex allowlist (categorized lightly).
        """
        skills = []
        found_skills = set()

        # Common programming languages and frameworks
        skill_patterns = [
            r"\b(Python|JavaScript|TypeScript|Java|C\+\+|C#|Go|Rust|Ruby|PHP|Swift|Kotlin|Scala|R|MATLAB)\b",
            r"\b(React|Angular|Vue|Next\.js|Nuxt|Django|Flask|FastAPI|Spring|Express\.js|Node\.js|\.NET)\b",
            r"\b(AWS|Azure|GCP|Docker|Kubernetes|Terraform|Ansible|Jenkins|Git|GitHub|GitLab|CI/CD)\b",
            r"\b(SQL|NoSQL|MongoDB|PostgreSQL|MySQL|Redis|Elasticsearch|GraphQL|REST)\b",
            r"\b(Machine Learning|Deep Learning|AI|Data Science|NLP|Computer Vision)\b",
        ]

        for pattern in skill_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                found_skills.add(match.group(1))

        for skill_name in found_skills:
            skills.append(
                Skill(
                    name=skill_name,
                    category=self._categorize_skill_name(skill_name),
                )
            )

        return skills

    def _split_date_range(self, date_range: str) -> tuple:
        """
        Split a date range string like 'Oct 2025 - Mar 2026' into start and end.

        Args:
            date_range: Date range string

        Returns:
            Tuple of (parsed_start_date_or_None, end_date_string_or_None)
        """
        if not date_range:
            return None, None

        # Common separators: dash, en-dash, em-dash, to
        for sep in [" – ", " - ", " — ", " to "]:
            if sep in date_range:
                parts = date_range.split(sep, 1)
                return self._parse_date(parts[0].strip()), parts[1].strip()

        # Single date (treat as start)
        return self._parse_date(date_range.strip()), None

    def _parse_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """
        Parse date string into datetime object.

        Args:
            date_str: Date string in various formats

        Returns:
            Datetime object or None
        """
        if not date_str:
            return None

        # Common date formats
        formats = [
            "%Y-%m",
            "%Y-%m-%d",
            "%m/%Y",
            "%m/%Y",
            "%b %Y",
            "%B %Y",
        ]

        for fmt in formats:
            try:
                return datetime.strptime(date_str.strip(), fmt)
            except ValueError:
                continue

        return None

    def _normalize_url(self, url: Optional[str]) -> Optional[str]:
        """
        Normalize URL by adding https:// scheme if missing.

        LLMs often return URLs without the scheme (e.g., 'linkedin.com/in/...')
        which causes pydantic HttpUrl validation to fail.

        Args:
            url: URL string that may or may not have a scheme

        Returns:
            Normalized URL with scheme, or None if input is None/empty
        """
        if not url or not isinstance(url, str):
            return None

        url = url.strip()

        # Already has a scheme
        if url.startswith(("http://", "https://")):
            return url

        # Add https:// scheme for common domains
        return f"https://{url}"
