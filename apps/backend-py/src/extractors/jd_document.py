"""
Job Raider - Job Description Document Extractor

Extract plain text from uploaded PDF or DOCX job descriptions.
Text only; no LLM and no profile writes.

Author: Job Raider
Date: 2026-08-20
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO, List, Union

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# Treat extracts shorter than this as near-empty (likely scanned / no text layer).
MIN_USEFUL_CHARS = 40

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx"})


@dataclass
class JdDocumentExtract:
    """Result of extracting text from a JD document.

    Attributes:
        text: Extracted plain text (may be empty).
        filename: Original upload filename.
        char_count: Character count of ``text`` after strip.
        warnings: Human-readable warnings (e.g. empty text layer).
    """

    text: str
    filename: str
    char_count: int
    warnings: List[str] = field(default_factory=list)


def _normalize_extension(filename: str) -> str:
    """Return the lowercased file extension including the leading dot.

    Args:
        filename: Original upload filename.

    Returns:
        Extension such as ``.pdf`` or ``.docx``, or empty string if none.
    """
    return Path(filename).suffix.lower()


def is_supported_jd_filename(filename: str) -> bool:
    """Return whether ``filename`` has a supported JD upload extension.

    Args:
        filename: Original upload filename.

    Returns:
        True for ``.pdf`` / ``.docx`` (case-insensitive); otherwise False.
    """
    return _normalize_extension(filename) in SUPPORTED_EXTENSIONS


def extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract plain text from PDF bytes via pypdf.

    Args:
        data: Raw PDF file contents.

    Returns:
        Concatenated page text. Empty string when no text layer is present.

    Raises:
        ValueError: If the PDF cannot be read.
    """
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: List[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            parts.append(page_text)
        return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f"Failed to extract text from PDF: {exc}") from exc


def extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract plain text from DOCX bytes via python-docx.

    Args:
        data: Raw DOCX file contents.

    Returns:
        Paragraph and table text joined with newlines.

    Raises:
        ValueError: If the DOCX cannot be read.
    """
    try:
        doc = Document(io.BytesIO(data))
        parts: List[str] = []
        for paragraph in doc.paragraphs:
            parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text for cell in row.cells]
                parts.append(" ".join(cell_texts))
        return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc


def extract_jd_document(
    data: bytes,
    filename: str,
) -> JdDocumentExtract:
    """Extract plain text from a JD PDF or DOCX upload.

    Does not invent content. Empty or near-empty extracts return a warning
    so the caller can ask the user to paste text instead (e.g. scanned PDF).

    Args:
        data: Raw file bytes.
        filename: Original filename used for extension detection and response.

    Returns:
        ``JdDocumentExtract`` with text, char count, and any warnings.

    Raises:
        ValueError: Unsupported extension or unreadable file.
    """
    ext = _normalize_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type. Upload a PDF or DOCX job description.")

    if ext == ".pdf":
        raw = extract_text_from_pdf_bytes(data)
    else:
        raw = extract_text_from_docx_bytes(data)

    text = raw.strip()
    warnings: List[str] = []
    if not text:
        warnings.append(
            "No extractable text found. The file may be a scanned PDF "
            "without a text layer. Paste the job description instead."
        )
    elif len(text) < MIN_USEFUL_CHARS:
        warnings.append(
            "Extracted text is very short. Check the file content, or paste "
            "the job description if the document is image-only."
        )

    logger.info(
        "JD document extract filename=%s chars=%s warnings=%s",
        filename,
        len(text),
        len(warnings),
    )
    return JdDocumentExtract(
        text=text,
        filename=filename,
        char_count=len(text),
        warnings=warnings,
    )


def extract_jd_document_from_stream(
    stream: Union[BinaryIO, bytes],
    filename: str,
) -> JdDocumentExtract:
    """Read bytes from a stream (or accept bytes) and extract JD text.

    Args:
        stream: File-like object with ``read()``, or raw bytes.
        filename: Original upload filename.

    Returns:
        ``JdDocumentExtract`` from :func:`extract_jd_document`.
    """
    if isinstance(stream, (bytes, bytearray)):
        data = bytes(stream)
    else:
        data = stream.read()
    return extract_jd_document(data, filename)
