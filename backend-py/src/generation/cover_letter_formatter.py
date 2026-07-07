"""
Job Raider - Cover Letter Formatter

Exports generated cover letters to DOCX and PDF. Much simpler than the
resume formatter because a cover letter is a single linear document.

Author: Job Raider
Date: 2026-06-29
"""

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, List, Optional

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from ..utils.logger import Components, get_logger

logger = get_logger(Components.GENERATION)


@dataclass
class CoverLetterExportOptions:
    """
    Options for cover-letter export.

    Attributes:
        sender_name: Name of the applicant (optional).
        sender_email: Email address (optional).
        sender_location: Location (optional).
        subject: Subject line; defaults to "Application for {title} at {company}".
        date: Date string; defaults to today's date.
        ats_mode: If True, produce plain, ATS-friendly output.
    """

    sender_name: Optional[str] = None
    sender_email: Optional[str] = None
    sender_location: Optional[str] = None
    subject: Optional[str] = None
    date: Optional[str] = None
    ats_mode: bool = False


@dataclass
class FormattedCoverLetter:
    """
    Result of a cover-letter export attempt.

    Attributes:
        pdf_path: Path to generated PDF, or None.
        docx_path: Path to generated DOCX, or None.
        success: Whether at least one format succeeded.
        errors: Human-readable errors from failed formats.
    """

    pdf_path: Optional[str] = None
    docx_path: Optional[str] = None
    success: bool = False
    errors: List[str] = field(default_factory=list)


class CoverLetterFormatter:
    """
    Format a cover letter as DOCX and/or PDF.

    The formatter is intentionally minimal: it adds an optional sender
    block, date, recipient, subject, and the letter body. If sender
    details are missing they are omitted rather than causing a failure.
    """

    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize the formatter.

        Args:
            output_dir: Directory for exported files. Defaults to
                ``data/outputs`` under the project root.
        """
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            base_dir = Path(__file__).resolve().parents[3]  # backend-py/
            self.output_dir = base_dir / "data" / "outputs"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def format_letter(
        self,
        content: str,
        filename: str,
        formats: List[str],
        company: str,
        title: str,
        options: Optional[CoverLetterExportOptions] = None,
    ) -> FormattedCoverLetter:
        """
        Export a cover letter to the requested formats.

        Args:
            content: Plain-text cover letter body.
            filename: Base filename without extension.
            formats: List of formats to generate, e.g. ``["docx", "pdf"]``.
            company: Company name for the recipient/subject lines.
            title: Job title for the subject line.
            options: Optional sender/subject/date overrides.

        Returns:
            ``FormattedCoverLetter`` with file paths and any errors.
        """
        if formats is None:
            formats = ["docx", "pdf"]
        if options is None:
            options = CoverLetterExportOptions()

        result = FormattedCoverLetter()

        if "docx" in formats:
            if DOCX_AVAILABLE:
                try:
                    result.docx_path = self._format_docx(
                        content, filename, company, title, options
                    )
                    result.success = True
                except Exception as exc:
                    message = f"DOCX generation failed: {exc}"
                    result.errors.append(message)
                    logger.error(message)
            else:
                message = "DOCX generation unavailable: python-docx not installed"
                result.errors.append(message)
                logger.warning(message)

        if "pdf" in formats:
            if REPORTLAB_AVAILABLE:
                try:
                    result.pdf_path = self._format_pdf(
                        content, filename, company, title, options
                    )
                    result.success = True
                except Exception as exc:
                    message = f"PDF generation failed: {exc}"
                    result.errors.append(message)
                    logger.error(message)
            else:
                message = "PDF generation unavailable: reportlab not installed"
                result.errors.append(message)
                logger.warning(message)

        return result

    def _build_subject(
        self, options: CoverLetterExportOptions, title: str, company: str
    ) -> str:
        """
        Build the subject line, falling back to a sensible default.

        Args:
            options: Export options, possibly containing a custom subject.
            title: Job title.
            company: Company name.

        Returns:
            Subject line string.
        """
        if options.subject:
            return options.subject
        return f"Application for {title} at {company}"

    def _format_docx(
        self,
        content: str,
        filename: str,
        company: str,
        title: str,
        options: CoverLetterExportOptions,
    ) -> str:
        """
        Format the cover letter as a DOCX file.

        Args:
            content: Cover letter body.
            filename: Base filename without extension.
            company: Company name.
            title: Job title.
            options: Export options.

        Returns:
            Path to the generated DOCX file.
        """
        filepath = self.output_dir / f"{filename}.docx"
        doc = Document()

        date_str = options.date or datetime.now().strftime("%B %d, %Y")
        subject = self._build_subject(options, title, company)

        # Sender block
        if options.sender_name:
            p = doc.add_paragraph()
            p.add_run(options.sender_name).bold = True
        if options.sender_email:
            doc.add_paragraph(options.sender_email)
        if options.sender_location:
            doc.add_paragraph(options.sender_location)

        doc.add_paragraph(date_str)

        # Recipient
        recipient = f"Hiring Manager\n{company}"
        doc.add_paragraph(recipient)

        # Subject
        subject_para = doc.add_paragraph()
        subject_para.add_run("Subject: ").bold = True
        subject_para.add_run(subject)

        # Body
        paragraphs = [p for p in content.split("\n\n") if p.strip()]
        for paragraph_text in paragraphs:
            doc.add_paragraph(paragraph_text.strip())

        doc.save(str(filepath))
        logger.info("Generated cover-letter DOCX: %s", filepath)
        return str(filepath)

    def _format_pdf(
        self,
        content: str,
        filename: str,
        company: str,
        title: str,
        options: CoverLetterExportOptions,
    ) -> str:
        """
        Format the cover letter as a PDF file.

        Args:
            content: Cover letter body.
            filename: Base filename without extension.
            company: Company name.
            title: Job title.
            options: Export options.

        Returns:
            Path to the generated PDF file.
        """
        filepath = self.output_dir / f"{filename}.pdf"
        date_str = options.date or datetime.now().strftime("%B %d, %Y")
        subject = self._build_subject(options, title, company)

        doc = SimpleDocTemplate(str(filepath), pagesize=LETTER)
        styles = getSampleStyleSheet()
        normal = styles["Normal"]
        normal.fontName = "Helvetica"
        normal.fontSize = 11
        normal.leading = 14

        subject_style = ParagraphStyle(
            "Subject",
            parent=normal,
            fontName="Helvetica-Bold",
            spaceAfter=12,
        )

        story: List[Any] = []

        if options.sender_name:
            story.append(Paragraph(f"<b>{options.sender_name}</b>", normal))
        if options.sender_email:
            story.append(Paragraph(options.sender_email, normal))
        if options.sender_location:
            story.append(Paragraph(options.sender_location, normal))
        if any([options.sender_name, options.sender_email, options.sender_location]):
            story.append(Spacer(1, 0.1 * inch))

        story.append(Paragraph(date_str, normal))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(f"Hiring Manager<br/>{company}", normal))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(f"<b>Subject:</b> {subject}", subject_style))
        story.append(Spacer(1, 0.15 * inch))

        paragraphs = [p for p in content.split("\n\n") if p.strip()]
        for paragraph_text in paragraphs:
            story.append(Paragraph(paragraph_text.strip(), normal))
            story.append(Spacer(1, 0.1 * inch))

        doc.build(story)
        logger.info("Generated cover-letter PDF: %s", filepath)
        return str(filepath)
