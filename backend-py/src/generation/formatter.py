"""
Job Raider - Resume Formatter

This module provides resume formatting functionality for PDF and DOCX output
with configurable templates, ATS-friendly mode, and section customization.

Author: Job Raider
Date: 2026-04-27
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
from dataclasses import dataclass, field
import io

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..utils.logger import get_logger, Components
from .resume_writer import GeneratedResume

logger = get_logger(Components.GENERATION)

# Default section order for resume generation
DEFAULT_SECTIONS_ORDER = [
    "summary", "skills", "experience", "projects", "education"
]


@dataclass
class FormatOptions:
    """
    Options for resume formatting.

    Controls template style, ATS mode, section visibility and ordering.

    Attributes:
        template: Template name (professional, modern, minimal, technical, executive)
        ats_mode: If True, produces ATS-friendly plain text formatting
        sections_order: Custom section ordering (default order if None)
        sections_hidden: Sections to omit from output
        sections_renamed: Map of section name to custom header text
    """
    template: str = "professional"
    ats_mode: bool = False
    sections_order: Optional[List[str]] = None
    sections_hidden: Optional[List[str]] = None
    sections_renamed: Optional[Dict[str, str]] = None


@dataclass
class FormattedResume:
    """
    Formatted resume with file paths.

    Attributes:
        pdf_path: Path to generated PDF file, or None
        docx_path: Path to generated DOCX file, or None
        success: Whether at least one format was generated
        errors: List of error messages from failed generations
    """
    pdf_path: Optional[str] = None
    docx_path: Optional[str] = None
    success: bool = False
    errors: List[str] = field(default_factory=list)


class TemplateManager:
    """
    Manage resume templates for different formats and styles.

    Provides pre-configured templates for common resume styles,
    each defining font, size, colors, and spacing.
    """

    TEMPLATES: Dict[str, Dict[str, Any]] = {
        "professional": {
            "font": "Times New Roman",
            "font_size": 11,
            "title_size": 18,
            "heading_size": 14,
            "heading_color": "#1a1a1a",
            "line_spacing": 1.15,
            "separator": "line",
        },
        "modern": {
            "font": "Arial",
            "font_size": 10,
            "title_size": 16,
            "heading_size": 13,
            "heading_color": "#2563eb",
            "line_spacing": 1.0,
            "separator": "line",
        },
        "minimal": {
            "font": "Helvetica",
            "font_size": 11,
            "title_size": 16,
            "heading_size": 13,
            "heading_color": "#333333",
            "line_spacing": 1.2,
            "separator": "space",
        },
        "technical": {
            "font": "Courier",
            "font_size": 10,
            "title_size": 14,
            "heading_size": 12,
            "heading_color": "#059669",
            "line_spacing": 1.0,
            "separator": "dash",
        },
        "executive": {
            "font": "Georgia",
            "font_size": 11,
            "title_size": 20,
            "heading_size": 14,
            "heading_color": "#1e3a5f",
            "line_spacing": 1.15,
            "separator": "double_line",
        },
    }

    @classmethod
    def get_template(cls, template_name: str = "professional") -> Dict[str, Any]:
        """
        Get template configuration by name.

        Args:
            template_name: Name of template to retrieve

        Returns:
            Template configuration dictionary with font, size, and color settings
        """
        return cls.TEMPLATES.get(template_name, cls.TEMPLATES["professional"])

    @classmethod
    def list_templates(cls) -> List[str]:
        """
        List all available template names.

        Returns:
            List of template name strings
        """
        return list(cls.TEMPLATES.keys())


class ResumeFormatter:
    """
    Format generated resumes into PDF and DOCX formats.

    Supports configurable templates with ATS-friendly mode,
    section reordering, and section customization.
    """

    def __init__(self, output_dir: str = "data/outputs"):
        """
        Initialize the resume formatter.

        Args:
            output_dir: Directory to save formatted resumes
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = get_logger(Components.GENERATION)

    def format_resume(
        self,
        resume: GeneratedResume,
        filename: Optional[str] = None,
        formats: Optional[List[str]] = None,
        options: Optional[FormatOptions] = None,
    ) -> FormattedResume:
        """
        Format a resume into PDF and/or DOCX.

        Args:
            resume: Generated resume to format
            filename: Base filename (without extension)
            formats: List of formats to generate ("pdf", "docx", or both)
            options: Formatting options including template, ATS mode, sections

        Returns:
            FormattedResume with file paths and status
        """
        if formats is None:
            formats = ["pdf", "docx"]
        if options is None:
            options = FormatOptions()

        if not self._get_ordered_sections(options):
            self.logger.warning("All resume sections hidden -- output will contain name only")

        result = FormattedResume()

        if not filename:
            timestamp = __import__('datetime').datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}"

        if "pdf" in formats and REPORTLAB_AVAILABLE:
            try:
                pdf_path = self._format_pdf(resume, filename, options)
                result.pdf_path = pdf_path
                result.success = True
            except Exception as e:
                result.errors.append(f"PDF generation failed: {str(e)}")
                self.logger.error(f"PDF generation failed: {str(e)}")

        if "docx" in formats and DOCX_AVAILABLE:
            try:
                docx_path = self._format_docx(resume, filename, options)
                result.docx_path = docx_path
                result.success = True
            except Exception as e:
                result.errors.append(f"DOCX generation failed: {str(e)}")
                self.logger.error(f"DOCX generation failed: {str(e)}")

        if result.pdf_path or result.docx_path:
            result.success = True

        return result

    def _get_section_header(
        self, section: str, options: FormatOptions
    ) -> str:
        """
        Get the display header for a section, applying renames.

        Args:
            section: Internal section name (e.g., "summary", "skills")
            options: Format options with potential rename mapping

        Returns:
            Display header text
        """
        headers = {
            "summary": "Professional Summary",
            "skills": "Skills",
            "experience": "Experience",
            "projects": "Projects",
            "education": "Education",
        }
        if options.sections_renamed and section in options.sections_renamed:
            return options.sections_renamed[section]
        return headers.get(section, section.title())

    def _get_ordered_sections(self, options: FormatOptions) -> List[str]:
        """
        Get sections in the configured order, excluding hidden ones.

        Args:
            options: Format options with ordering and hiding config

        Returns:
            Ordered list of section names to render
        """
        hidden = set(options.sections_hidden or [])
        order = options.sections_order or DEFAULT_SECTIONS_ORDER
        valid_sections = set(DEFAULT_SECTIONS_ORDER)

        # Filter to known sections, warn on typos
        result = [s for s in order if s not in hidden]
        for section in result:
            if section not in valid_sections:
                self.logger.warning("Unknown section '%s' in sections_order, skipping", section)
        return [s for s in result if s in valid_sections]

    def _has_section_data(self, resume: GeneratedResume, section: str) -> bool:
        """
        Check if a section has data to render.

        Args:
            resume: Generated resume data
            section: Section name to check

        Returns:
            True if section has non-empty data
        """
        data_map = {
            "summary": bool(resume.summary),
            "skills": bool(resume.skills),
            "experience": bool(resume.experience),
            "projects": bool(resume.projects),
            "education": bool(resume.education),
        }
        return data_map.get(section, False)

    def _format_pdf(
        self,
        resume: GeneratedResume,
        filename: str,
        options: FormatOptions,
    ) -> str:
        """
        Format resume as PDF with template-aware styling.

        Args:
            resume: Generated resume data
            filename: Base filename without extension
            options: Format options controlling template and layout

        Returns:
            Path to generated PDF file
        """
        filepath = self.output_dir / f"{filename}.pdf"
        template = TemplateManager.get_template(options.template)

        doc = SimpleDocTemplate(str(filepath), pagesize=LETTER)
        base_styles = getSampleStyleSheet()

        # Build styles from template config
        heading_color = colors.HexColor(template["heading_color"])

        if options.ats_mode:
            title_style = ParagraphStyle(
                'ATSTitle', parent=base_styles['Normal'],
                fontSize=14, spaceAfter=6,
            )
            heading_style = ParagraphStyle(
                'ATSHeading', parent=base_styles['Normal'],
                fontSize=12, spaceAfter=4, spaceBefore=10,
            )
        else:
            title_style = ParagraphStyle(
                'CustomTitle', parent=base_styles['Heading1'],
                fontSize=template["title_size"],
                textColor=heading_color, spaceAfter=12,
            )
            heading_style = ParagraphStyle(
                'CustomHeading', parent=base_styles['Heading2'],
                fontSize=template["heading_size"],
                textColor=heading_color, spaceAfter=8, spaceBefore=12,
            )

        normal_style = base_styles['Normal']
        bullet_char = "-" if options.ats_mode else "•"

        story: List[Any] = []

        # Name/Title from summary
        lines = resume.summary.split('\n') if resume.summary else []
        if lines:
            name = lines[0].split('-')[0].strip()
            story.append(Paragraph(name, title_style))
            story.append(Spacer(0.1 * inch))

        # Render sections in configured order
        for section in self._get_ordered_sections(options):
            if not self._has_section_data(resume, section):
                continue

            header_text = self._get_section_header(section, options)

            # Section heading
            if options.ats_mode:
                story.append(Paragraph(header_text.upper(), heading_style))
            else:
                story.append(Paragraph(header_text, heading_style))

            # Section separator
            if not options.ats_mode:
                sep = template.get("separator", "space")
                if sep == "line":
                    story.append(HRFlowable(
                        width="100%", thickness=0.5,
                        color=heading_color, spaceAfter=4,
                    ))
                elif sep == "dash":
                    story.append(HRFlowable(
                        width="100%", thickness=0.5,
                        color=colors.grey, spaceAfter=4, dash=[3, 3],
                    ))
                elif sep == "double_line":
                    story.append(HRFlowable(
                        width="100%", thickness=1.0,
                        color=heading_color, spaceAfter=2,
                    ))
                    story.append(HRFlowable(
                        width="100%", thickness=0.5,
                        color=heading_color, spaceAfter=4,
                    ))

            # Section content
            if section == "summary":
                story.append(Paragraph(resume.summary, normal_style))

            elif section == "skills":
                skills_text = ", ".join(resume.skills)
                story.append(Paragraph(skills_text, normal_style))

            elif section == "experience":
                for exp in resume.experience:
                    title_line = f"{exp.get('title', 'Unknown')} at {exp.get('company', 'Unknown')}"
                    story.append(Paragraph(title_line, base_styles['Heading3']))
                    story.append(Paragraph(exp.get('dates', ''), normal_style))
                    for hl in exp.get('highlights', []):
                        story.append(Paragraph(f"{bullet_char} {hl}", normal_style))
                    story.append(Spacer(0.1 * inch))

            elif section == "projects":
                for project in resume.projects:
                    story.append(Paragraph(project['name'], base_styles['Heading3']))
                    if project.get('description'):
                        story.append(Paragraph(project['description'], normal_style))
                    if project.get('technologies'):
                        tech_text = ", ".join(project['technologies'])
                        story.append(Paragraph(
                            f"<b>Technologies:</b> {tech_text}", normal_style
                        ))
                    for hl in project.get('highlights', []):
                        story.append(Paragraph(f"{bullet_char} {hl}", normal_style))
                    story.append(Spacer(0.1 * inch))

            elif section == "education":
                for edu in resume.education:
                    edu_line = edu['degree']
                    if edu.get('school'):
                        edu_line += f" - {edu['school']}"
                    if edu.get('year'):
                        edu_line += f" ({edu['year']})"
                    story.append(Paragraph(edu_line, normal_style))

            story.append(Spacer(0.15 * inch))

        doc.build(story)
        self.logger.info("Generated PDF: %s", filepath)
        return str(filepath)

    def _format_docx(
        self,
        resume: GeneratedResume,
        filename: str,
        options: FormatOptions,
    ) -> str:
        """
        Format resume as DOCX with template-aware styling.

        Args:
            resume: Generated resume data
            filename: Base filename without extension
            options: Format options controlling template and layout

        Returns:
            Path to generated DOCX file
        """
        filepath = self.output_dir / f"{filename}.docx"
        template = TemplateManager.get_template(options.template)

        doc = Document()
        heading_color_hex = template["heading_color"].lstrip('#')
        heading_rgb = RGBColor.from_string(heading_color_hex)
        bullet_char = "-" if options.ats_mode else "•"

        # Name/Title from summary
        lines = resume.summary.split('\n') if resume.summary else []
        if lines:
            name = lines[0].split('-')[0].strip()
            heading = doc.add_heading(name, 0)
            if not options.ats_mode:
                for run in heading.runs:
                    run.font.color.rgb = heading_rgb

        # Render sections in configured order
        for section in self._get_ordered_sections(options):
            if not self._has_section_data(resume, section):
                continue

            header_text = self._get_section_header(section, options)

            # Section heading
            if options.ats_mode:
                h = doc.add_heading(header_text.upper(), level=1)
            else:
                h = doc.add_heading(header_text, level=1)
                for run in h.runs:
                    run.font.color.rgb = heading_rgb

            # Section content
            if section == "summary":
                doc.add_paragraph(resume.summary)

            elif section == "skills":
                p = doc.add_paragraph()
                for i, skill in enumerate(resume.skills):
                    p.add_run(skill)
                    if i < len(resume.skills) - 1:
                        p.add_run(', ')

            elif section == "experience":
                for exp in resume.experience:
                    p = doc.add_paragraph()
                    p.add_run(exp.get('title', 'Unknown'), bold=True)
                    p.add_run(f" at {exp.get('company', 'Unknown')}")
                    p.add_run(f"\n{exp.get('dates', '')}", italic=True)

                    for hl in exp.get('highlights', []):
                        doc.add_paragraph(hl, style='List Bullet')

            elif section == "projects":
                for project in resume.projects:
                    doc.add_heading(project['name'], level=2)

                    if project.get('description'):
                        doc.add_paragraph(project['description'])

                    if project.get('technologies'):
                        p = doc.add_paragraph()
                        p.add_run('Technologies: ', bold=True)
                        for i, tech in enumerate(project['technologies']):
                            p.add_run(tech)
                            if i < len(project['technologies']) - 1:
                                p.add_run(', ')

                    for hl in project.get('highlights', []):
                        doc.add_paragraph(hl, style='List Bullet')

            elif section == "education":
                for edu in resume.education:
                    p = doc.add_paragraph()
                    p.add_run(edu['degree'], bold=True)
                    if edu.get('school'):
                        p.add_run(f" - {edu['school']}")
                    if edu.get('year'):
                        p.add_run(f" ({edu['year']})")

        doc.save(str(filepath))
        self.logger.info("Generated DOCX: %s", filepath)
        return str(filepath)
